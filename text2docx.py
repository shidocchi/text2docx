import io
import os
import re
import sys
import argparse
from typing import Iterator
from docx import Document
from docx.shared import Mm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

__version__ = '0.1.5'

class Text2Docx:
  """text typesetter"""

  PAGESEP = '\x0C'

  PAGE = {
    'a3': (297, 420),
    'b4': (257, 364),
    'a4': (210, 297),
    'b5': (182, 257),
    'a5': (148, 210),
    'hagaki': (100, 148),
  }

  FONT = {
    'lc': 'Lucida Console',
    'lst': 'Lucida Sans Typewriter',
  }

  EAFONT = {
    'biz': 'BIZ UDゴシック',
    'hg': 'HGｺﾞｼｯｸ',
    'hge': 'HGｺﾞｼｯｸE',
    'hgm': 'HGｺﾞｼｯｸM',
    'meiryo': 'メイリオ',
    'yu': '游ゴシック',
    'ms': 'ＭＳ ゴシック',
  }

  HEAD_NUMBER = ' [Page {PAGE}/{NUMPAGES}]'
  head_parser = re.compile(r'(\{\w+\})|([^{}]*)')
  HEAD_ALIGN = {
    (False, False): WD_ALIGN_PARAGRAPH.CENTER,
    (True, False):  WD_ALIGN_PARAGRAPH.RIGHT,
    (False, True):  WD_ALIGN_PARAGRAPH.LEFT,
    (True, True):   WD_ALIGN_PARAGRAPH.CENTER,
  }

  SAMPLE = [
    'The quick brown fox jumps over the lazy dog',
    '色は匂へど散りぬるを我が世誰ぞ常ならむ有為の奥山今日越えて浅き夢見し酔ひもせず',
  ]

  def __init__(self, textin) -> None:
    self.args = self.get_args()
    if not self.args.raw:
      textin = io.TextIOWrapper(textin.buffer, encoding='utf-8')
    self.doc = Document()
    self.set_section(self.doc.sections[0])
    self.set_style(self.doc.styles['Normal'])
    if self.args.col:
      self.set_multicolumn(self.doc.sections[0], self.args.col)
    if self.args.sample:
      self.set_sample()
    else:
      self.typeset(textin)

  def get_args(self) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
      prog='python -m text2docx',
      description='text typesetter')
    parser.add_argument('--raw', help='suppress stdin encoding',
      action='store_true')
    parser.add_argument('--out', help='output filename')
    parser.add_argument('--page', help='page size',
      choices=self.PAGE.keys())
    parser.add_argument('--landscape', help='landscape',
      action='store_true')
    parser.add_argument('--margin', help='margin mm',
      type=float,
      nargs=4, metavar=('top','bottom','left','right'))
    parser.add_argument('--col', help='multi column',
      type=int,
      choices=(2,3))
    parser.add_argument('--size', help='font pt',
      type=float)
    parser.add_argument('--font', help='font',
      choices=self.FONT.keys())
    parser.add_argument('--eafont', help='eastasia font',
      choices=self.EAFONT.keys())
    parser.add_argument('--sample', help='font sample',
      action='store_true')
    parser.add_argument('--do', help='operation',
      choices=['print', 'edit', 'open'])
    head_args = parser.add_mutually_exclusive_group()
    head_args.add_argument('--number', help='page number on header',
      action='store_true')
    head_args.add_argument('--header', help='header')
    parser.add_argument('--footer', help='footer')
    parser.set_defaults(out='output.docx')
    parser.set_defaults(page='a4', margin=(10,10,10,10))
    parser.set_defaults(size=14, font='lc', eafont='hge')
    return parser.parse_args()

  def set_section(self, sect) -> None:
    if self.args.landscape:
      sect.orientation = WD_ORIENT.LANDSCAPE
      (sect.page_height,
       sect.page_width) = map(Mm, self.PAGE[self.args.page])
    else:
      sect.orientation = WD_ORIENT.PORTRAIT
      (sect.page_width,
       sect.page_height) = map(Mm, self.PAGE[self.args.page])
    (sect.top_margin,
     sect.bottom_margin,
     sect.left_margin,
     sect.right_margin) = map(Mm, self.args.margin)
    (sect.header_distance,
     sect.footer_distance) = map(Mm, [5, 5])
    if self.args.number:
      self.set_head(sect.header, self.HEAD_NUMBER)
    elif self.args.header:
      self.set_head(sect.header, self.args.header)
    if self.args.footer:
      self.set_head(sect.footer, self.args.footer)

  def set_multicolumn(self, sect, num) -> None:
    sectPr = sect._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), str(num))

  def set_style(self, sty) -> None:
    sty.font.size = Pt(self.args.size)
    sty.font.name = self.FONT.get(self.args.font, self.args.font)
    sty.element.rPr.rFonts.set(qn('w:eastAsia'),
      self.EAFONT.get(self.args.eafont, self.args.eafont))

  def save(self) -> None:
    self.doc.save(self.args.out)
    if self.args.do:
      os.startfile(self.args.out, operation=self.args.do)

  def typeset(self, textin, sep=PAGESEP) -> None:
    for page in self.paginate(textin, sep):
      if page == sep:
        self.doc.add_page_break()
      else:
        self.doc.add_paragraph(page)

  def paginate(self, textin, sep) -> Iterator[str]:
    page = []
    for line in textin:
      while True:
        part = line.partition(sep)
        page.append(part[0])
        if part[1] == '':
          break
        else:
          yield ''.join(page)
          yield sep
          page = []
          line = part[2]
          if not line.rstrip():
            break
    if page:
      yield ''.join(page)

  def set_head(self, head, fcode) -> None:
    par = head.paragraphs[0]
    par.alignment = self.HEAD_ALIGN[(fcode.startswith(' '), fcode.endswith(' '))]
    for m in self.head_parser.finditer(fcode):
      if m.group(1):
        self.add_field(par, m.group(1)[1:-1])
      else:
        par.add_run(m.group(2))

  def add_field(self, par, text) -> None:
    run = par.add_run()
    run._r.append(OxmlElement('w:fldChar'))
    run._r[-1].set(qn('w:fldCharType'), 'begin')
    run._r.append(OxmlElement('w:instrText'))
    run._r[-1].text = text
    run._r.append(OxmlElement('w:fldChar'))
    run._r[-1].set(qn('w:fldCharType'), 'end')

  def set_sample(self):
    table = self.doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = 'font name'
    hdr[1].text = ''
    for k,fn in self.FONT.items():
      row = table.add_row().cells
      row[0].width = Mm(50)
      row[1].width = Mm(150)
      row[0].text = '{0} ({1})'.format(fn,k)
      r = row[1].paragraphs[0].add_run(self.SAMPLE[0])
      r.font.name = fn
    for k,fn in self.EAFONT.items():
      row = table.add_row().cells
      row[0].width = Mm(50)
      row[1].width = Mm(150)
      row[0].text = '{0} ({1})'.format(fn,k)
      r = row[1].paragraphs[0].add_run(self.SAMPLE[1])
      r.font.name = self.FONT['lc']
      r._element.rPr.rFonts.set(qn('w:eastAsia'), fn)

if __name__ == '__main__':
  d = Text2Docx(sys.stdin)
  d.save()
